"""DOCX readers: specification or narrative from Word documents."""

from pathlib import Path
from typing import Any

from ..utils import get_logger

_log = get_logger(__name__)


def read_docx_spec(path: str | Path) -> dict[str, Any] | None:
    """Read a DOCX specification (e.g. Activities Threshold Logic Matrix.docx).

    Returns structured text/paragraphs for inclusion in the report.
    Uses python-docx if available; otherwise returns a minimal placeholder.

    Args:
        path: Path to the .docx file.

    Returns:
        Dict with keys such as 'paragraphs', 'tables', or None if unreadable.
    """
    path = Path(path)
    if not path.exists():
        _log.warning("DOCX spec not found: %s", path)
        return None
    try:
        import docx
        doc = docx.Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = []
        for t in doc.tables:
            rows = [[c.text for c in row.cells] for row in t.rows]
            tables.append(rows)
        return {"paragraphs": paragraphs, "tables": tables}
    except ImportError:
        _log.warning("python-docx not installed; DOCX content will be skipped")
        return {"paragraphs": [], "tables": [], "error": "python-docx not installed"}
    except Exception as e:
        _log.exception("Failed to read DOCX %s: %s", path, e)
        return None
